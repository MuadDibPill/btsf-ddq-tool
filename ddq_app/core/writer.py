"""
DDQ Automation — Word Document Output
Generates a formatted .docx DDQ from a list of Answer objects.
Matches the BTSF IB/PE house style.
"""

import os
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import (OUTPUT_DIR, COLOUR_ANSWERED, COLOUR_PARTIAL,
                    COLOUR_GAP, COLOUR_NA, COLOUR_NAVY, COLOUR_STEEL,
                    COLOUR_GOLD, COLOUR_LGRAY)
from core.generator import Answer
from core.questions import Question


# ── Colour helpers ────────────────────────────────────────────────────────────

def _rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour.upper())
    tcPr.append(shd)


def _set_para_border_bottom(para, hex_colour: str, size: int = 6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_colour.upper())
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Paragraph / run helpers ───────────────────────────────────────────────────

def _heading_para(doc: Document, text: str):
    """Navy full-width section header paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    # Shading via XML
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  COLOUR_NAVY)
    pPr.append(shd)
    run = p.add_run(f"  {text}")
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = _rgb("FFFFFF")
    run.font.name = "Arial"
    return p


def _q_label(para, qid: str):
    run = para.add_run(f"{qid}  ")
    run.bold           = True
    run.font.size      = Pt(9)
    run.font.color.rgb = _rgb(COLOUR_STEEL)
    run.font.name      = "Arial"


def _q_text(para, text: str):
    run = para.add_run(text)
    run.bold           = True
    run.font.size      = Pt(9)
    run.font.color.rgb = _rgb("1A1A1A")
    run.font.name      = "Arial"


def _answer_run(para, text: str, colour: str):
    run = para.add_run(text)
    run.font.size      = Pt(9)
    run.font.color.rgb = _rgb(colour)
    run.font.name      = "Arial"
    return run


def _badge(para, label: str, colour: str):
    run = para.add_run(f"  [{label}]")
    run.bold           = True
    run.font.size      = Pt(8)
    run.font.color.rgb = _rgb(colour)
    run.font.name      = "Arial"


def _source_run(para, sources: List[str]):
    if not sources:
        return
    run = para.add_run("  Sources: " + ", ".join(sources))
    run.italic         = True
    run.font.size      = Pt(7.5)
    run.font.color.rgb = _rgb("5A6472")
    run.font.name      = "Arial"


# ── Table helpers ─────────────────────────────────────────────────────────────

def _kv_table(doc: Document, rows: List[tuple]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    for i, (label, value) in enumerate(rows):
        bg = COLOUR_LGRAY if i % 2 == 0 else "FFFFFF"
        c0, c1 = table.rows[i].cells
        _set_cell_bg(c0, bg)
        _set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(8.5); r0.font.name = "Arial"
        r0.font.color.rgb = _rgb(COLOUR_NAVY)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(8.5); r1.font.name = "Arial"
    return table


def _signals_table(doc: Document, signals: Dict[str, bool]):
    active   = [k for k, v in signals.items() if v]
    inactive = [k for k, v in signals.items() if not v]

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Deal-type signals detected")
    run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
    run.font.color.rgb = _rgb(COLOUR_NAVY)

    table = doc.add_table(rows=1 + len(signals), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.5)

    # Header
    h0, h1 = table.rows[0].cells
    for cell, txt in [(h0, "Signal"), (h1, "Detected")]:
        _set_cell_bg(cell, COLOUR_NAVY)
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(8.5); r.font.name = "Arial"
        r.font.color.rgb = _rgb("FFFFFF")

    for i, (k, v) in enumerate(signals.items(), 1):
        c0, c1 = table.rows[i].cells
        bg = COLOUR_LGRAY if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(c0, bg); _set_cell_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(k.replace("_", " ").title())
        r0.font.size = Pt(8.5); r0.font.name = "Arial"
        r1 = c1.paragraphs[0].add_run("Yes ✓" if v else "No")
        r1.font.size = Pt(8.5); r1.font.name = "Arial"
        r1.font.color.rgb = _rgb(COLOUR_ANSWERED if v else "5A6472")


def _summary_table(doc: Document, answers: List[Answer]):
    counts = {"answered": 0, "partial": 0, "gap": 0}
    for a in answers:
        counts[a.confidence] = counts.get(a.confidence, 0) + 1
    total = len(answers)

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    w = Inches(1.6)
    for col in table.columns:
        col.width = w

    labels = ["Total questions", "Answered", "Partial", "Gap / Missing"]
    values = [str(total),
              f"{counts['answered']} ({counts['answered']*100//total}%)" if total else "0",
              f"{counts['partial']} ({counts['partial']*100//total}%)"  if total else "0",
              f"{counts['gap']}     ({counts['gap']*100//total}%)"      if total else "0"]
    colours = [COLOUR_NAVY, COLOUR_ANSWERED, COLOUR_PARTIAL, COLOUR_GAP]

    for i, (label, value, colour) in enumerate(zip(labels, values, colours)):
        c0 = table.rows[0].cells[i]
        c1 = table.rows[1].cells[i]
        _set_cell_bg(c0, colour)
        _set_cell_bg(c1, COLOUR_LGRAY)
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(8); r0.font.name = "Arial"
        r0.font.color.rgb = _rgb("FFFFFF")
        r1 = c1.paragraphs[0].add_run(value)
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Arial"
        r1.font.color.rgb = _rgb(colour)


# ── Main export function ──────────────────────────────────────────────────────

CONFIDENCE_COLOUR = {
    "answered": COLOUR_ANSWERED,
    "partial":  COLOUR_PARTIAL,
    "gap":      COLOUR_GAP,
    "n/a":      COLOUR_NA,
}

CONFIDENCE_LABEL = {
    "answered": "Answered",
    "partial":  "Partial",
    "gap":      "Gap — pending",
    "n/a":      "N/A",
}


def export_docx(answers: List[Answer],
                signals: Dict[str, bool],
                deal_name: str = "Project XXX",
                site_info: Dict[str, str] = None) -> str:
    """
    Generate a filled DDQ Word document.
    Returns the path to the saved file.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Cover ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run(deal_name.upper())
    r.bold = True; r.font.size = Pt(22); r.font.name = "Arial"
    r.font.color.rgb = _rgb(COLOUR_NAVY)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("BTSF Due Diligence Questionnaire — Auto-generated")
    r2.font.size = Pt(11); r2.font.name = "Arial"
    r2.font.color.rgb = _rgb(COLOUR_STEEL)

    # Gold rule
    p3 = doc.add_paragraph()
    _set_para_border_bottom(p3, COLOUR_GOLD, 8)

    # Timestamp + confidence note
    p4 = doc.add_paragraph()
    r4 = p4.add_run(
        f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}  |  "
        "Items marked [Gap — pending] must be completed prior to DDQ submission."
    )
    r4.italic = True; r4.font.size = Pt(8.5); r4.font.name = "Arial"
    r4.font.color.rgb = _rgb("5A6472")

    # Site info table
    if site_info:
        doc.add_paragraph()
        _kv_table(doc, list(site_info.items()))

    # Signals table
    doc.add_paragraph()
    _signals_table(doc, signals)

    # Summary scorecard
    doc.add_paragraph()
    p_sc = doc.add_paragraph()
    r_sc = p_sc.add_run("Completion summary")
    r_sc.bold = True; r_sc.font.size = Pt(10); r_sc.font.name = "Arial"
    r_sc.font.color.rgb = _rgb(COLOUR_NAVY)
    _summary_table(doc, answers)

    doc.add_page_break()

    # ── Legend ────────────────────────────────────────────────────────────────
    for colour, label in [
        (COLOUR_ANSWERED, "Answered — sourced from data room"),
        (COLOUR_PARTIAL,  "Partial — incomplete; gap noted"),
        (COLOUR_GAP,      "Gap — not found in data room; must be provided"),
    ]:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(f"  {label}")
        lr.font.size = Pt(8); lr.font.name = "Arial"
        lr.font.color.rgb = _rgb(colour)

    doc.add_paragraph()

    # ── Q&A sections ─────────────────────────────────────────────────────────
    current_section = None

    for ans in answers:
        # Section heading
        if ans.question.section != current_section:
            current_section = ans.question.section
            _heading_para(doc, current_section.upper())

        # Question paragraph
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after  = Pt(2)
        _q_label(qp, ans.question.qid)
        _q_text(qp, ans.question.text)

        # Answer paragraph
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(4)

        colour = CONFIDENCE_COLOUR.get(ans.confidence, COLOUR_GAP)
        label  = CONFIDENCE_LABEL.get(ans.confidence, "Gap — pending")

        if ans.confidence == "gap":
            _answer_run(ap, "[No information found in data room — item must be provided.]",
                        COLOUR_GAP)
            if ans.gap_note:
                _answer_run(ap, f" {ans.gap_note}", COLOUR_GAP)
        else:
            _answer_run(ap, ans.text, "1A1A1A")

        _badge(ap, label, colour)

        # Gap note for partial answers
        if ans.confidence == "partial" and ans.gap_note:
            gp = doc.add_paragraph()
            gp.paragraph_format.left_indent = Inches(0.25)
            gp.paragraph_format.space_after = Pt(2)
            gr = gp.add_run(f"Missing: {ans.gap_note}")
            gr.italic = True; gr.font.size = Pt(8); gr.font.name = "Arial"
            gr.font.color.rgb = _rgb(COLOUR_PARTIAL)

        # Sources footnote
        if ans.sources:
            sp2 = doc.add_paragraph()
            sp2.paragraph_format.left_indent = Inches(0.25)
            sp2.paragraph_format.space_after = Pt(6)
            _source_run(sp2, ans.sources)

    # ── Appendix — document list ──────────────────────────────────────────────
    doc.add_page_break()
    _heading_para(doc, "APPENDIX — KEY DOCUMENTS IN DATA ROOM")

    # Collect unique sources
    all_sources = sorted(set(
        s for ans in answers for s in ans.sources
    ))
    if all_sources:
        for src in all_sources:
            sp3 = doc.add_paragraph()
            sp3.paragraph_format.space_after = Pt(2)
            sr = sp3.add_run(f"  {src}")
            sr.font.size = Pt(8.5); sr.font.name = "Arial"
    else:
        sp3 = doc.add_paragraph()
        sp3.add_run("  No source documents found.").font.size = Pt(8.5)

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    dp = doc.add_paragraph()
    dr = dp.add_run(
        "DISCLAIMER — This document is confidential and prepared solely for BTSF's due diligence "
        "process. All financial projections are illustrative and subject to material change. "
        "This document does not constitute an offer to lend or invest. Items marked "
        "[Gap — pending] represent open items to be resolved prior to closing."
    )
    dr.italic = True; dr.font.size = Pt(8); dr.font.name = "Arial"
    dr.font.color.rgb = _rgb("5A6472")

    # ── Save ──────────────────────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = deal_name.replace(" ", "_").replace("/", "-")
    filename  = f"DDQ_{safe_name}_{timestamp}.docx"
    filepath  = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    print(f"\n[Output] Saved: {filepath}")
    return filepath
